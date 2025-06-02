import os
from docx import Document

# Define the source directory and the consolidated output file
source_dir = "SAS/"
consolidated_output_file = "all_sas_documents_content.txt"

# Check if the source directory exists
if not os.path.exists(source_dir):
    print(f"EROARE: Directorul sursă '{source_dir}' nu a fost găsit!")
    exit()

print(f"Se procesează fișierele din directorul '{source_dir}'...")
found_docx_files = False
all_extracted_text_content = [] # List to hold all text before writing to file

# Clear the consolidated file if it already exists, to start fresh
if os.path.exists(consolidated_output_file):
    try:
        os.remove(consolidated_output_file)
        print(f"Fișierul existent '{consolidated_output_file}' a fost șters pentru a crea unul nou.")
    except Exception as e:
        print(f"Nu s-a putut șterge fișierul existent '{consolidated_output_file}': {e}")
        # Decide if you want to exit() or try to append; for safety, let's exit if we can't ensure a fresh start.
        exit()

for filename in os.listdir(source_dir):
    if filename.endswith(".docx"):
        found_docx_files = True
        source_filepath = os.path.join(source_dir, filename)
        
        all_extracted_text_content.append(f"\n--- START OF DOCUMENT: {filename} ---\n")
        try:
            print(f"  Se extrage textul din '{source_filepath}'...")
            doc = Document(source_filepath)
            
            for para in doc.paragraphs:
                all_extracted_text_content.append(para.text + "\n")
            
            if doc.tables:
                all_extracted_text_content.append("\n--- TABLES IN DOCUMENT ---\n")
                for i, table in enumerate(doc.tables):
                    all_extracted_text_content.append(f"--- Table {i+1} ---\n")
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text_parts = []
                            for para_in_cell in cell.paragraphs:
                                cell_text_parts.append(para_in_cell.text)
                            row_text.append(" | ".join(cell_text_parts)) # Join cell paragraphs, then join cells
                        all_extracted_text_content.append(" || ".join(row_text) + "\n") # Join cells with ' || '
                    all_extracted_text_content.append("--- End of Table ---\n")
            
            all_extracted_text_content.append(f"--- END OF DOCUMENT: {filename} ---\n\n")
            print(f"    Textul din '{filename}' a fost adăugat pentru consolidare.")

        except Exception as e:
            all_extracted_text_content.append(f"### EROARE la procesarea fișierului '{source_filepath}': {e} ###\n")
            print(f"    EROARE la procesarea fișierului '{source_filepath}': {e}")

if not found_docx_files:
    print(f"Nu s-au găsit fișiere .docx în directorul '{source_dir}'.")
else:
    try:
        with open(consolidated_output_file, "w", encoding="utf-8") as f:
            f.writelines(all_extracted_text_content) # Use writelines as all_extracted_text_content contains strings ending in \n
        print(f"\nTot textul extras a fost consolidat în '{consolidated_output_file}'.")
    except Exception as e:
        print(f"\nEROARE la scrierea fișierului consolidat '{consolidated_output_file}': {e}")

print("Procesul de extracție consolidată a textului a fost finalizat.") 